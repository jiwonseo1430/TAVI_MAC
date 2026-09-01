# 12_manuscript_docx.py — convert manuscript/manuscript.md to Word.
# Writes <data_dir>/output/manuscript.docx (binary outputs live on OneDrive).
# Minimal markdown support: #/##/### headings, **bold** inline, bullet lists,
# --- rules (skipped), plain paragraphs. Run from repo root.

import os
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


def data_dir():
    if os.environ.get("TAVI_MAC_DATA"):
        return os.environ["TAVI_MAC_DATA"]
    with open(os.path.join("config", "config.yml"), encoding="utf-8") as f:
        m = re.search(r'data_dir:\s*"?([^"\r\n]+)"?', f.read())
    if not m:
        raise SystemExit("data_dir not found in config/config.yml")
    return m.group(1).strip()


SRC = os.path.join("manuscript", "manuscript.md")
OUT = os.path.join(data_dir(), "output", "manuscript.docx")

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)


def add_runs(p, text):
    """Split **bold** and *italic* segments into runs."""
    for part in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            p.add_run(part[1:-1]).italic = True
        elif part:
            p.add_run(part)


def para(text, bullet=False):
    p = doc.add_paragraph(style="List Bullet" if bullet else None)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)
    return p


def heading(text, level):
    p = doc.add_paragraph()
    run_text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    r = p.add_run(run_text)
    r.bold = True
    r.font.size = Pt({1: 14, 2: 13, 3: 12}.get(level, 12))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


with open(SRC, encoding="utf-8") as f:
    lines = f.read().splitlines()

buf = []


def flush():
    global buf
    if buf:
        para(" ".join(buf))
        buf = []


for line in lines:
    s = line.rstrip()
    if s.startswith("#"):
        flush()
        level = len(s) - len(s.lstrip("#"))
        heading(s.lstrip("#").strip(), level)
    elif s.strip() == "---":
        flush()
    elif s.startswith("- "):
        flush()
        para(s[2:].strip(), bullet=True)
    elif re.match(r"^\d+[\.–-]", s.strip()) and "References" not in s:
        flush()
        para(s.strip())
    elif s.strip() == "":
        flush()
    else:
        buf.append(s.strip())
flush()

doc.save(OUT)
print("written:", OUT)
