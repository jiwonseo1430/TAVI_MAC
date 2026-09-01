# 11_tables_docx.py — assemble manuscript Tables 1-4 into one Word file.
# Reads CSVs from <data_dir>/output/, writes <data_dir>/output/Tables_1-4.docx
# Data dir: env TAVI_MAC_DATA, else config/config.yml (data_dir key).
# Run from repo root: python analysis/scripts/11_tables_docx.py

import csv
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def data_dir():
    if os.environ.get("TAVI_MAC_DATA"):
        return os.environ["TAVI_MAC_DATA"]
    with open(os.path.join("config", "config.yml"), encoding="utf-8") as f:
        m = re.search(r'data_dir:\s*"?([^"\r\n]+)"?', f.read())
    if not m:
        raise SystemExit("data_dir not found in config/config.yml")
    return m.group(1).strip()


OUT = os.path.join(data_dir(), "output")


def read_csv(name):
    with open(os.path.join(OUT, name), encoding="utf-8-sig", newline="") as f:
        return [row for row in csv.reader(f)]


LABELS = {
    "log2_mac": "MAC volume (per doubling)",
    "any_mac": "MAC (vs no MAC)",
    "mac_group_medLow": "Low MAC (0 < volume ≤ 108.6 mm³)",
    "mac_group_medHigh": "High MAC (volume > 108.6 mm³)",
    "sex_female": "Sex (female)",
    "bmi": "Body mass index",
    "sts": "STS score (%)",
    "ckd": "Chronic kidney disease",
    "afib": "Atrial fibrillation",
    "ldl_10": "LDL cholesterol (per 10 mg/dL)",
}


def hline(cell, edge):
    """Add a single horizontal border (edge = 'top'/'bottom') to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "6")
    el.set(qn("w:color"), "000000")
    borders.append(el)


def style_cell(cell, text, bold=False, center=False, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold


def acad_table(doc, header, rows, widths_cm, bold_rows=(), indent_rows=()):
    """Three-line academic table: rules above/below header and below last row."""
    t = doc.add_table(rows=len(rows) + 1, cols=len(header))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        style_cell(c, h, bold=True, center=(j > 0))
        hline(c, "top")
        hline(c, "bottom")
    for i, r in enumerate(rows):
        for j, v in enumerate(r):
            c = t.rows[i + 1].cells[j]
            txt = ("    " + v) if (j == 0 and i in indent_rows) else v
            style_cell(c, txt, bold=(i in bold_rows), center=(j > 0))
            if i == len(rows) - 1:
                hline(c, "bottom")
    for j, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[j].width = Cm(w)
    return t


def title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)


def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(8)
    run.font.name = "Calibri"
    p.paragraph_format.space_before = Pt(6)


doc = Document()

# ---------- Table 1 ----------
rows = read_csv("table1_binary.csv")
title(doc, "Table 1. Baseline characteristics by MAC presence")
acad_table(doc, [h.replace("p-value", "P value") for h in rows[0]], rows[1:],
           [4.2, 3.0, 3.0, 3.0, 2.0])
note(doc, "Values are mean ± SD or n (%). MAC = mitral annular calcification; "
     "BMI = body mass index; STS = Society of Thoracic Surgeons; HTN = hypertension; "
     "DM = diabetes mellitus; CKD = chronic kidney disease; COPD = chronic obstructive "
     "pulmonary disease; PAD = peripheral artery disease; CAD = coronary artery disease; "
     "A.fib = atrial fibrillation; MI = myocardial infarction; TC = total cholesterol; "
     "TG = triglycerides; LDL/HDL = low-/high-density lipoprotein cholesterol; "
     "LVEF = left ventricular ejection fraction; MDPG = mean diastolic pressure gradient "
     "of the mitral valve; RSVP = right ventricular systolic pressure; MR = mitral regurgitation.")
doc.add_page_break()

# ---------- Table 2 ----------
rows = read_csv("table2.csv")[1:]


def t2_label(t):
    s = t.strip()
    if s == "Low":
        return "Low MAC (≤108.6 mm³)"
    if s == "High":
        return "High MAC (>108.6 mm³)"
    return t


indent = {i for i, r in enumerate(rows) if r[0].startswith("  ")}
title(doc, "Table 2. Univariable Cox regression for all-cause mortality (N = 525, 102 deaths)")
acad_table(doc, ["Variable", "HR (95% CI)", "P value"],
           [[t2_label(r[0]), r[1], r[2]] for r in rows],
           [6.0, 4.0, 2.5], indent_rows=indent)
note(doc, "HR = hazard ratio; CI = confidence interval. Abbreviations as in Table 1.")
doc.add_page_break()

# ---------- Table 3 ----------
body, bold_rows, indent_rows = [], set(), set()
for hdr, f in [("Model 1 — Continuous MAC volume (primary)", "table3c_final.csv"),
               ("Model 2 — MAC presence", "table3c_binary.csv"),
               ("Model 3 — MAC volume category", "table3c_categories.csv")]:
    bold_rows.add(len(body))
    body.append([hdr, "", ""])
    for r in read_csv(f)[1:]:
        indent_rows.add(len(body))
        body.append([LABELS.get(r[0], r[0]), r[1], r[2]])
title(doc, "Table 3. Multivariable Cox models for all-cause mortality (N = 525, 102 deaths)")
acad_table(doc, ["Variable", "HR (95% CI)", "P value"], body,
           [7.0, 4.0, 2.5], bold_rows=bold_rows, indent_rows=indent_rows)
note(doc, "Covariates were selected by backward elimination (univariable p<0.10 entry, "
     "retention p<0.05) with the MAC exposure forced into each model. "
     "Abbreviations as in Tables 1–2.")
doc.add_page_break()

# ---------- Table 4 ----------
rows = read_csv("table_time_interval.csv")[1:]
body, bold_rows, indent_rows = [], set(), set()
current = None
for r in rows:
    if r[0] != current:
        current = r[0]
        bold_rows.add(len(body))
        body.append([r[0], "", "", r[4]])
    indent_rows.add(len(body))
    body.append([r[1].replace("0-1 y", "0–1 year").replace(">1 y", ">1 year"),
                 r[2], r[3], ""])
title(doc, "Table 4. Period-specific adjusted hazard ratios for the MAC exposure")
acad_table(doc, ["Exposure / period", "HR (95% CI)", "P value", "Interaction P"],
           body, [5.5, 3.6, 2.2, 2.6], bold_rows=bold_rows, indent_rows=indent_rows)
note(doc, "Counting-process Cox models split at 1 year (365.25 days), adjusted for sex, "
     "body mass index, STS score, chronic kidney disease, atrial fibrillation, and LDL "
     "cholesterol (held constant across periods). Interaction P: Wald test of the "
     "difference between period-specific coefficients.")

out = os.path.join(OUT, "Tables_1-4.docx")
doc.save(out)
print("written:", out)
